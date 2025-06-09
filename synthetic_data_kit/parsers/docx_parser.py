# Copyright (c) Meta Platforms, Inc. and affiliates.
# All rights reserved.
#
# This source code is licensed under the terms described in the LICENSE file in
# the root directory of this source tree.
# DOCX parsers
import os
import logging
from pathlib import Path
import io
from xml.etree import ElementTree as ET

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

class DOCXParser:
    """Parser for Microsoft Word documents"""

    def _extract_image(self, element, doc):
        """Extract image from a drawing element."""
        try:
            # Define the namespaces
            namespaces = {
                'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
                'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'
            }
            
            # Find the blip element with proper namespace
            blip = element.find('.//a:blip', namespaces)
            if blip is not None:
                # Get the image ID from the r:embed attribute
                image_id = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                if image_id:
                    image_part = doc.part.related_parts[image_id]
                    return image_part.blob
        except Exception as e:
            logger.warning(f"Failed to extract image: {str(e)}")
        return None

    def parse(self, file_path: str, multimodal: bool = False) -> any:
        """Parse a DOCX file.
        
        Args:
            file_path: Path to the DOCX file
            multimodal: If True, extract text and images. Otherwise, extract text only.
            
        Returns:
            If multimodal is False, returns a string with extracted text.
            If multimodal is True, returns a list of dictionaries, where each dictionary
            has 'text' and 'image' keys.
            
        Raises:
            ValueError: If file cannot be read or is empty.
        """
        try:
            import docx
        except ImportError:
            raise ImportError("python-docx is required for DOCX parsing. Install with: pip install python-docx")

        logger.info(f"Starting to parse: {file_path}")
        
        try:
            doc = docx.Document(file_path)
            if not doc.paragraphs and not doc.tables:
                raise ValueError("DOCX file appears to be empty")
            
            logger.info(f"Found {len(doc.paragraphs)} paragraphs and {len(doc.tables)} tables")
            
            if not multimodal:
                # Text-only mode: combine all content
                text_content = []
                
                # Extract text from paragraphs
                for para in doc.paragraphs:
                    if para.text.strip():
                        text_content.append(para.text)
                
                # Extract text from tables
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if cell.text.strip():
                                text_content.append(cell.text)
                
                return "\n".join(text_content)
            else:
                # Multimodal mode: process content paragraph by paragraph
                output_data = []
                current_block = []
                current_image = None
                image_count = 0
                
                # Process paragraphs
                for para in doc.paragraphs:
                    if not para.text.strip():
                        continue
                        
                    # Check for images in the paragraph
                    para_image = None
                    for run in para.runs:
                        for element in run._element.iter():
                            if element.tag.endswith('drawing'):
                                para_image = self._extract_image(element, doc)
                                if para_image:
                                    image_count += 1
                                    logger.info(f"Found image {image_count} in paragraph")
                                    break
                        if para_image:
                            break
                    
                    # If we found an image, save the current block and start a new one
                    if para_image:
                        if current_block:
                            output_data.append({
                                'text': "\n".join(current_block),
                                'image': current_image
                            })
                            current_block = []
                        current_image = para_image
                    
                    # Add the paragraph text
                    current_block.append(para.text)
                
                # Process tables
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if cell.text.strip():
                                current_block.append(cell.text)
                
                # Add any remaining content
                if current_block:
                    output_data.append({
                        'text': "\n".join(current_block),
                        'image': current_image
                    })
                
                logger.info(f"Extracted {len(output_data)} content blocks with {image_count} images")
                return output_data
                
        except Exception as e:
            logger.error(f"Error processing DOCX: {str(e)}")
            raise

    def save(self, content: any, output_path: str) -> None:
        """Save the extracted content to a Lance file.
        
        Args:
            content: Extracted content (string or list of dicts)
            output_path: Path to save the Lance file
        """
        logger.info(f"Saving content to {output_path}...")
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        
        import lance
        import pyarrow as pa
        
        if isinstance(content, str):
            # Text-only mode
            data = [pa.array([content])]
            names = ['text']
            table = pa.Table.from_arrays(data, names=names)
        elif isinstance(content, list):
            # Multimodal mode
            texts = pa.array([item['text'] for item in content], type=pa.string())
            images_data = []
            for item in content:
                img = item.get('image')
                images_data.append(img if img is not None else None)
            images = pa.array(images_data, type=pa.binary())
            
            schema = pa.schema([
                ('text', pa.string()),
                ('image', pa.binary())
            ])
            table = pa.Table.from_arrays([texts, images], schema=schema)
        else:
            raise ValueError("Unsupported content type for saving.")
            
        lance.write_dataset(table, output_path, mode="overwrite")
        logger.info("Save completed successfully")